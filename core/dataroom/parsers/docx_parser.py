"""
Word document parser using python-docx for .docx files.
"""

from pathlib import Path
from .base import BaseParser, ParseResult


class DocxParser(BaseParser):
    """Extract text and tables from Word documents (.docx, .doc)."""

    extensions = [".docx", ".doc"]

    def parse(self, filepath: str) -> ParseResult:
        """Extract text and tables from a Word document.

        For .docx files, uses python-docx to extract:
        - Paragraphs (with heading detection for section structure)
        - Tables (as {headers, rows} dicts)
        - Core document properties as metadata

        For .doc files, returns an error since python-docx does not support
        the legacy binary format.

        Returns:
            ParseResult with full text, tables, and document metadata.
        """
        path = Path(filepath)
        if not path.exists():
            return ParseResult(error=f"File not found: {filepath}")

        if path.suffix.lower() == ".doc":
            return ParseResult(
                error="Legacy .doc format not supported. Convert to .docx first."
            )

        try:
            from docx import Document
        except ImportError:
            return ParseResult(
                error="python-docx not installed. Run: pip install python-docx"
            )

        try:
            doc = Document(str(path))
        except Exception as e:
            return ParseResult(error=f"Failed to open DOCX: {e}")

        # Extract paragraphs
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings and add markdown-style formatting
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                text_parts.append(f"\n# {text}")
            elif "heading 2" in style_name:
                text_parts.append(f"\n## {text}")
            elif "heading 3" in style_name:
                text_parts.append(f"\n### {text}")
            elif "title" in style_name:
                text_parts.append(f"\n# {text}")
            else:
                text_parts.append(text)

        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            if not table.rows:
                continue

            # First row as headers
            headers = [cell.text.strip() for cell in table.rows[0].cells]

            rows = []
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            tables.append({
                "headers": headers,
                "rows": rows,
                "table_index": i,
            })

        full_text = "\n".join(text_parts)

        # Extract metadata from core properties
        metadata = {"filename": path.name}
        try:
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.title:
                metadata["title"] = props.title
            if props.created:
                metadata["created"] = props.created.isoformat()[:10]
            if props.modified:
                metadata["modified"] = props.modified.isoformat()[:10]
        except Exception:
            pass

        # Estimate page count from paragraph count (rough: ~40 paragraphs per page)
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        estimated_pages = max(1, para_count // 40 + 1)

        return ParseResult(
            text=full_text,
            tables=tables,
            metadata=metadata,
            page_count=estimated_pages,
        )
